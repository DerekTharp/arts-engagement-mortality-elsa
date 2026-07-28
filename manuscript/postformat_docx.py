"""
Post-format a pandoc-generated submission .docx so tables and paragraphs look
clean in Word, without modifying the shared reference.docx template.

Fixes:
  - table cells single-spaced (pandoc's Compact style inherits the body's
    double spacing, which balloons row heights)
  - sensible column widths for the two main-manuscript tables (wider label
    column so model/variable names don't wrap to 2-3 lines)
  - booktabs horizontal rules (top, under header, bottom)
  - rows kept together (no splitting across a page break) and header row
    repeated at the top of each page
  - first-line indent restored on body paragraphs (Body Text style)

All table-property and cell-property elements are inserted in the order the
OOXML schema (CT_TblPr / CT_TcPr / CT_TrPr) requires. Word tolerates
out-of-order children, but stricter renderers (LibreOffice, Apple) compute
fixed-layout column widths from a schema-ordered tree and collapse the grid
otherwise.

Usage: python3 postformat_docx.py <file.docx> [more.docx ...]
"""

import os
import sys
import zipfile

import lxml.etree as ET
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ANS = "http://schemas.openxmlformats.org/drawingml/2006/main"

# Styles whose monospace font is intentional (code / syntax highlighting); the
# font sweep leaves these alone so any verbatim text stays fixed-width.
MONO_STYLES = {"macro", "Macro Text Char", "Source Code", "Verbatim Char",
               "HTML Preformatted", "HTML Preformatted Char"}

# Canonical successor tags for each managed element: the element is inserted
# immediately before the first of these that already exists in the parent
# (else appended). Lists follow the ECMA-376 sequence for the parent type.
TBLPR_AFTER = {
    "w:tblW": ("w:tblJc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders",
               "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
               "w:tblCaption", "w:tblDescription"),
    "w:tblBorders": ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
                     "w:tblCaption", "w:tblDescription"),
    "w:tblLayout": ("w:tblCellMar", "w:tblLook", "w:tblCaption",
                    "w:tblDescription"),
}
TCPR_AFTER = {
    "w:tcW": ("w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders", "w:shd",
              "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
              "w:vAlign", "w:hideMark"),
    "w:tcBorders": ("w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
                    "w:tcFitText", "w:vAlign", "w:hideMark"),
}
TRPR_AFTER = {
    "w:cantSplit": ("w:trHeight", "w:tblHeader", "w:tblCellSpacing", "w:jc",
                    "w:hidden"),
    "w:tblHeader": ("w:tblCellSpacing", "w:jc", "w:hidden"),
}

# Full ECMA-376 child sequence for each properties element. A final reorder
# pass sorts existing children into this order, so the script is idempotent:
# re-running it on an already-processed file still yields a schema-valid tree.
TBLPR_ORDER = ("w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
               "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW",
               "w:tblJc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders",
               "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
               "w:tblCaption", "w:tblDescription")
TCPR_ORDER = ("w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge",
              "w:tcBorders", "w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
              "w:tcFitText", "w:vAlign", "w:hideMark")
TRPR_ORDER = ("w:cnfStyle", "w:divId", "w:gridBefore", "w:gridAfter",
              "w:wBefore", "w:wAfter", "w:cantSplit", "w:trHeight",
              "w:tblHeader", "w:tblCellSpacing", "w:jc", "w:hidden")


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def _border(tag, val, sz=None):
    e = _el("w:" + tag)
    e.set(qn("w:val"), val)
    if val != "nil":
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
    return e


def _put(parent, tag, after_tags):
    """Remove any existing <tag> children, then insert a fresh one in
    schema-correct position. Returns the new element for the caller to fill."""
    for x in parent.findall(qn(tag)):
        parent.remove(x)
    e = OxmlElement(tag)
    parent.insert_element_before(e, *after_tags)
    return e


def _reorder(parent, order):
    """Sort the parent's children into the canonical schema sequence. Stable,
    so any element not in `order` keeps its relative position at the end."""
    rank = {t: i for i, t in enumerate(order)}
    kids = sorted(list(parent),
                  key=lambda c: rank.get("w:" + c.tag.split("}")[-1], len(order)))
    for c in kids:
        parent.append(c)


def add_rules(table, grid=False):
    """Borders for the table. Default is booktabs: 1pt top and bottom of the
    table, 0.5pt under the header row, no vertical or inter-row lines. With
    grid=True, every cell is boxed with a thin rule and the header row carries a
    heavier bottom rule — readable for a dense, many-row checklist."""
    tblPr = table._tbl.tblPr
    tb = _put(tblPr, "w:tblBorders", TBLPR_AFTER["w:tblBorders"])
    if grid:
        for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
            tb.append(_border(edge, "single", 4))
    else:
        tb.append(_border("top", "single", 8))
        tb.append(_border("bottom", "single", 8))
        for edge in ("left", "right", "insideH", "insideV"):
            tb.append(_border(edge, "nil"))
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcb = _put(tcPr, "w:tcBorders", TCPR_AFTER["w:tcBorders"])
        tcb.append(_border("bottom", "single", 8 if grid else 4))


def _single_space(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    sp = _el("w:spacing", line="240", lineRule="auto", before="0", after="40")
    pstyle = pPr.find(qn("w:pStyle"))
    if pstyle is not None:
        pstyle.addnext(sp)
    else:
        pPr.insert(0, sp)


def fix_table(
    table,
    widths=None,
    repeat_header=True,
    grid=False,
    font_size=None,
    cell_margin=None,
):
    tbl = table._tbl
    tblPr = tbl.tblPr
    add_rules(table, grid=grid)
    _put(tblPr, "w:tblLayout", TBLPR_AFTER["w:tblLayout"]).set(qn("w:type"), "fixed")
    if widths:
        tw = _put(tblPr, "w:tblW", TBLPR_AFTER["w:tblW"])
        tw.set(qn("w:w"), str(sum(widths)))
        tw.set(qn("w:type"), "dxa")
        grid = tbl.find(qn("w:tblGrid"))
        for gc in list(grid):
            grid.remove(gc)
        for wd in widths:
            grid.append(_el("w:gridCol", w=wd))
    for ri, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        if not trPr.findall(qn("w:cantSplit")):
            _put(trPr, "w:cantSplit", TRPR_AFTER["w:cantSplit"])
        if ri == 0 and repeat_header and not trPr.findall(qn("w:tblHeader")):
            _put(trPr, "w:tblHeader", TRPR_AFTER["w:tblHeader"])
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            if widths and ci < len(widths):
                tcW = _put(tcPr, "w:tcW", TCPR_AFTER["w:tcW"])
                tcW.set(qn("w:w"), str(widths[ci]))
                tcW.set(qn("w:type"), "dxa")
            if cell_margin is not None:
                tcMar = _put(
                    tcPr,
                    "w:tcMar",
                    ("w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"),
                )
                for edge in ("top", "left", "bottom", "right"):
                    margin = _el(
                        f"w:{edge}",
                        w=0 if edge in ("top", "bottom") else cell_margin,
                        type="dxa",
                    )
                    tcMar.append(margin)
            for p in cell.paragraphs:
                _single_space(p)
                if font_size is not None:
                    for run in p.runs:
                        run.font.size = Pt(font_size)
            _reorder(tcPr, TCPR_ORDER)
        _reorder(trPr, TRPR_ORDER)
    _reorder(tblPr, TBLPR_ORDER)


def _scaled_widths(available, proportions):
    """Convert proportions to twips while preserving the exact total width."""
    if abs(sum(proportions) - 1.0) > 1e-9:
        raise ValueError(f"Column-width proportions must sum to 1: {proportions}")
    widths = [int(available * p) for p in proportions[:-1]]
    widths.append(available - sum(widths))
    return widths


def _table_profile(filename, headers, ncol, available):
    """Return document-aware column widths, grid setting, and font size.

    Pandoc's automatic widths are especially poor for the dense supplement and
    STROBE tables. Profiles are keyed by semantic headers rather than table
    position so a newly inserted table cannot silently inherit the wrong
    geometry.
    """
    first = headers[0]
    proportions = None
    grid = False
    font_size = None
    cell_margin = None

    if filename == "manuscript.docx":
        font_size = 9.5
        cell_margin = 40
        if first == "Step (exposure handling)" and ncol == 5:
            proportions = (0.31, 0.19, 0.16, 0.205, 0.135)
        elif first == "Current-exposure contrast" and ncol == 6:
            proportions = (0.18, 0.17, 0.10, 0.10, 0.225, 0.225)
            font_size = 8.5
            cell_margin = 25

    elif filename == "strobe_checklist.docx" and first == "Item No" and ncol == 4:
        # The location column needs enough room for section references; the
        # section-heading rows are merged across all four columns below.
        proportions = (0.10, 0.22, 0.25, 0.43)
        grid = True
        font_size = 8.5
        cell_margin = 30

    elif filename == "supplement.docx":
        font_size = 9
        cell_margin = 40
        if first == "Statistic" and ncol == 2:
            proportions = (0.68, 0.32)
        elif first == "Time-varying confounder" and ncol == 3:
            proportions = (0.45, 0.25, 0.30)
        elif first == "Confounder" and ncol == 7:
            proportions = (0.28,) + (0.12,) * 6
            font_size = 8
            cell_margin = 20
        elif first == "Weight truncation" and ncol == 3:
            proportions = (0.34, 0.48, 0.18)
        elif first == "Variable" and ncol == 5:
            proportions = (0.36, 0.16, 0.16, 0.16, 0.16)
        elif first == "Model" and ncol == 5:
            if headers[1].startswith("Frequent vs never"):
                proportions = (0.28, 0.245, 0.245, 0.115, 0.115)
            else:
                proportions = (0.25, 0.255, 0.255, 0.12, 0.12)
        elif first == "From" and ncol == 4:
            proportions = (0.25, 0.25, 0.25, 0.25)
        elif first == "PGS" and ncol == 3:
            proportions = (0.36, 0.44, 0.20)
        elif first == "Covariate" and ncol == 3:
            proportions = (0.38, 0.30, 0.32)

    widths = _scaled_widths(available, proportions) if proportions else None
    return widths, grid, font_size, cell_margin


def enforce_font(path, font="Times New Roman"):
    """Force `font` across the whole document. Pandoc's reference document keeps
    Calibri/Cambria theme fonts and points the document default and the heading
    styles at those theme fonts, so any text not explicitly styled Times New
    Roman (headings, the title, theme-defaulted runs) renders in the wrong face.

    Three coordinated changes close every fallback path:
      1. rewrite the theme major/minor fonts to `font`
      2. give the document an explicit default font instead of a theme reference
      3. replace theme font references on every (non-code) style with `font`

    Operates on the .docx zip directly because the theme part is not exposed by
    python-docx."""
    def W(t):
        return f"{{{WNS}}}{t}"

    def A(t):
        return f"{{{ANS}}}{t}"

    mono_faces = {"Courier", "Courier New", "Consolas", "Monaco", "Lucida Console"}

    def force_rfonts(rf):
        for attr in list(rf.attrib):
            local = attr.split("}")[-1]
            # *Theme refs use mixed casing in the schema (asciiTheme but cstheme)
            if local.lower().endswith("theme") or local in ("ascii", "hAnsi", "cs", "eastAsia"):
                del rf.attrib[attr]
        for slot in ("ascii", "hAnsi", "cs", "eastAsia"):
            rf.set(W(slot), font)

    def force_run(r):
        """Force the font on a single run, preserving any monospace face."""
        rpr = r.find(W("rPr"))
        if rpr is None:
            rpr = ET.Element(W("rPr"))
            r.insert(0, rpr)
        rf = rpr.find(W("rFonts"))
        if rf is None:
            rf = ET.Element(W("rFonts"))
            rstyle = rpr.find(W("rStyle"))
            rstyle.addnext(rf) if rstyle is not None else rpr.insert(0, rf)
        if rf.get(W("ascii")) in mono_faces:
            return
        force_rfonts(rf)

    with zipfile.ZipFile(path) as zin:
        order = zin.namelist()
        items = {n: zin.read(n) for n in order}

    theme = "word/theme/theme1.xml"
    if theme in items:
        th = ET.fromstring(items[theme])
        for group in ("majorFont", "minorFont"):
            g = th.find(f".//{A(group)}")
            latin = g.find(A("latin")) if g is not None else None
            if latin is not None:
                latin.set("typeface", font)
        items[theme] = ET.tostring(th, xml_declaration=True, encoding="UTF-8", standalone=True)

    styles = "word/styles.xml"
    st = ET.fromstring(items[styles])
    dd = st.find(W("docDefaults"))
    if dd is not None:
        rprd = dd.find(W("rPrDefault"))
        if rprd is None:
            rprd = ET.SubElement(dd, W("rPrDefault"))
        rpr = rprd.find(W("rPr"))
        if rpr is None:
            rpr = ET.SubElement(rprd, W("rPr"))
        rf = rpr.find(W("rFonts"))
        if rf is None:
            rf = ET.Element(W("rFonts"))
            rpr.insert(0, rf)
        force_rfonts(rf)
    for style in st.findall(W("style")):
        name_el = style.find(W("name"))
        name = name_el.get(W("val")) if name_el is not None else style.get(W("styleId"))
        if name in MONO_STYLES:
            continue
        for rf in style.iter(W("rFonts")):  # top-level rPr and nested tblStylePr
            force_rfonts(rf)
    items[styles] = ET.tostring(st, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Run-level sweep over body, headers and footers: pandoc occasionally emits
    # theme-referenced runs (e.g. inline elements) that override the style font.
    for part in order:
        base = part.rsplit("/", 1)[-1]
        if part == "word/document.xml" or (
            part.startswith("word/") and base.endswith(".xml")
            and (base.startswith("header") or base.startswith("footer"))
        ):
            tree = ET.fromstring(items[part])
            for r in tree.iter(W("r")):
                force_run(r)
            items[part] = ET.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    tmp = path + ".fonttmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in order:
            zout.writestr(n, items[n])
    os.replace(tmp, path)


def main(path):
    doc = Document(path)
    filename = os.path.basename(path)
    sec = doc.sections[0]
    # Some reference.docx templates carry no explicit page geometry (python-docx
    # then returns None); default to US Letter with 1-inch margins.
    if sec.page_width is None:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin = sec.bottom_margin = Inches(1)
    avail = int((sec.page_width - sec.left_margin - sec.right_margin) / 635)  # twips
    for table in doc.tables:
        headers = tuple(cell.text.strip() for cell in table.rows[0].cells)
        hdr0 = headers[0]
        ncol = len(table.columns)
        widths, grid, font_size, cell_margin = _table_profile(
            filename, headers, ncol, avail
        )
        fix_table(
            table,
            widths,
            grid=grid,
            font_size=font_size,
            cell_margin=cell_margin,
        )
        if hdr0 == "Item No" and ncol == 4:
            # Section-header rows carry text only in column 0 (e.g. "Methods").
            # Merge the row so the heading spans the full width instead of
            # wrapping in the narrow first column.
            for row in table.rows[1:]:
                cells = row.cells
                if cells[0].text.strip() and not any(c.text.strip() for c in cells[1:]):
                    merged = cells[0]
                    for c in cells[1:]:
                        merged = merged.merge(c)
    # First-line indent on narrative body paragraphs only. Pandoc styles the
    # first paragraph after a heading "First Paragraph" (stays flush) and the
    # rest "Body Text" (indented). We then zero the indent on front-matter
    # (title page, before the first heading) and on bold-led paragraphs (the
    # structured-abstract, key-messages, and Declarations labels), so a label
    # like "Methods" or "Data sharing:" sits at the margin.
    for st in doc.styles:
        if st.name in ("Body Text", "BodyText"):
            st.paragraph_format.first_line_indent = Inches(0.3)
    seen_heading = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            seen_heading = True
            continue
        if (
            p.runs
            and p.runs[0].bold
            and p.text.strip().startswith(("Table ", "Table S", "Figure ", "Figure S"))
        ):
            p.paragraph_format.keep_with_next = True
        if p.style.name in ("Body Text", "BodyText"):
            first_bold = bool(p.runs and p.runs[0].bold)
            if (not seen_heading) or first_bold:
                p.paragraph_format.first_line_indent = Inches(0)
    doc.save(path)
    enforce_font(path)
    print(f"  formatted {path}  ({len(doc.tables)} tables, usable width {avail} twips)")


if __name__ == "__main__":
    for p in sys.argv[1:]:
        main(p)
